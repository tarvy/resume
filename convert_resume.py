#!/usr/bin/env python3
"""
Render resume outputs (DOCX, HTML, MD, PDF) from a YAML/JSON manifest.
"""

import json
import re
import html
from copy import deepcopy

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

MANIFEST_PATH = 'resume.yaml'
HTML_TEMPLATE_PATH = 'templates/resume.html.tmpl'
HTML_EMBED_TEMPLATE_PATH = 'templates/resume.embed.tmpl'
DOCX_PATH = 'TravisGlassResume.docx'
HTML_PATH = 'TravisGlassResume.html'
MD_PATH = 'TravisGlassResume.md'
PDF_PATH = 'TravisGlassResume.pdf'

ORDINAL_RE = re.compile(r'(\d+)(st|nd|rd|th)\b')
URL_RE = re.compile(r'(https?://[^\s]+|www\.[^\s]+)')


def make_url(text):
    """Ensure URL has a protocol prefix."""
    if text.startswith('www.'):
        return 'https://' + text
    return text


def load_manifest(path):
    """Load YAML if available; fallback to JSON (YAML 1.2 JSON-compatible)."""
    with open(path, 'r', encoding='utf-8') as f:
        content = f.read().strip()
    try:
        import yaml  # type: ignore
        return yaml.safe_load(content)
    except Exception:
        try:
            return json.loads(content)
        except json.JSONDecodeError as exc:
            raise RuntimeError(
                'Manifest is not JSON-compatible YAML. Install PyYAML '
                'or keep resume.yaml in JSON-style YAML syntax.'
            ) from exc


def normalize_date(date_text):
    return re.sub(r'\s-\s', ' \u2013 ', date_text.strip())


def set_style_fonts(style, font_name, size_pt):
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    rfonts = style._element.rPr.rFonts
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)
    rfonts.set(qn('w:eastAsia'), font_name)
    rfonts.set(qn('w:cs'), font_name)


def set_table_indent(table, inches):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tbl_pr)
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), str(int(inches * 1440)))
    tbl_ind.set(qn('w:type'), 'dxa')


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    for edge, edge_data in kwargs.items():
        edge_tag = qn(f'w:{edge}')
        edge_el = tc_borders.find(edge_tag)
        if edge_el is None:
            edge_el = OxmlElement(f'w:{edge}')
            tc_borders.append(edge_el)
        for key, value in edge_data.items():
            edge_el.set(qn(f'w:{key}'), str(value))


def clear_cell(cell):
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn('w:p'):
            tc.remove(child)


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    r_pr.append(color)

    # Underline
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    r_pr.append(underline)

    new_run.append(r_pr)

    text_el = OxmlElement('w:t')
    text_el.text = text
    new_run.append(text_el)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def apply_paragraph_format(paragraph, line_spacing=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = line_spacing


def add_text_runs(paragraph, text, bold=False, italic=False):
    lines = text.split('\n')
    for line_idx, line in enumerate(lines):
        if line_idx > 0:
            paragraph.add_run().add_break()
        segments = line.split('\t')
        for seg_idx, segment in enumerate(segments):
            last = 0
            for match in ORDINAL_RE.finditer(segment):
                prefix = segment[last:match.start()]
                if prefix:
                    run = paragraph.add_run(prefix)
                    run.bold = bold
                    run.italic = italic
                run = paragraph.add_run(match.group(1))
                run.bold = bold
                run.italic = italic
                sup = paragraph.add_run(match.group(2))
                sup.bold = bold
                sup.italic = italic
                sup.font.superscript = True
                last = match.end()
            tail = segment[last:]
            if tail:
                run = paragraph.add_run(tail)
                run.bold = bold
                run.italic = italic
            if seg_idx < len(segments) - 1:
                paragraph.add_run().add_tab()


def render_html_text(text, linkify=False):
    def render_segment(segment):
        parts = []
        last = 0
        for match in ORDINAL_RE.finditer(segment):
            parts.append(html.escape(segment[last:match.start()], quote=False))
            parts.append(html.escape(match.group(1), quote=False))
            parts.append(f'<sup>{html.escape(match.group(2), quote=False)}</sup>')
            last = match.end()
        parts.append(html.escape(segment[last:], quote=False))
        return ''.join(parts)

    def linkify_text(text):
        parts = []
        last = 0
        for match in URL_RE.finditer(text):
            parts.append(html.escape(text[last:match.start()], quote=False))
            url_text = match.group(1)
            url_href = make_url(url_text)
            parts.append(f'<a href="{html.escape(url_href)}">{html.escape(url_text, quote=False)}</a>')
            last = match.end()
        parts.append(html.escape(text[last:], quote=False))
        return ''.join(parts)

    lines = text.split('\n')
    rendered_lines = []
    for line in lines:
        segments = line.split('\t')
        if linkify:
            rendered_segments = [linkify_text(seg) for seg in segments]
        else:
            rendered_segments = [render_segment(seg) for seg in segments]
        rendered_lines.append('<span class="tab"></span>'.join(rendered_segments))
    return '<br/>'.join(rendered_lines)


def load_template(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()


def render_html(data, include_wrapper=True):
    template_path = HTML_TEMPLATE_PATH if include_wrapper else HTML_EMBED_TEMPLATE_PATH
    template = load_template(template_path)

    name_row = '\n'.join([
        '<tr>',
        '<td></td>',
        f'<td class="name-cell" colspan="2">{html.escape(data["name"], quote=False)}</td>',
        '</tr>',
    ])

    contact_lines = '\n'.join(
        f'<div>{render_html_text(line, linkify=True)}</div>' for line in data['contact_lines']
    )
    contact_row = '\n'.join([
        '<tr class="contact-row">',
        '<td></td>',
        '<td class="contact-cell" colspan="2">',
        contact_lines,
        '</td>',
        '</tr>',
    ])

    edu = data['education']
    education_row = '\n'.join([
        '<tr class="section-row">',
        '<td></td>',
        '<td><strong>EDUCATION</strong></td>',
        '<td>',
        f'<p><strong>{html.escape(edu["degree"], quote=False)}</strong></p>',
        f'<p>{html.escape(edu["institution"], quote=False)}</p>',
        f'<p><strong>Major:</strong> {html.escape(edu["major"], quote=False)}</p>',
        f'<p><strong>Specialization:</strong> {html.escape(edu["specialization"], quote=False)}</p>',
        '</td>',
        '</tr>',
    ])

    experience_header_row = '\n'.join([
        '<tr class="section-row">',
        '<td></td>',
        '<td><em><strong>EXPERIENCE</strong></em></td>',
        '<td></td>',
        '</tr>',
    ])

    experience_rows = []
    for job in data['experience']:
        date_html = html.escape(normalize_date(job['date']), quote=False)
        title = html.escape(job['title'], quote=False)
        company = html.escape(job['company'], quote=False)
        row_parts = [
            '<tr class="job-row">',
            '<td></td>',
            f'<td><p><em>{date_html}</em></p></td>',
            '<td>',
            f'<p><strong>{title}, </strong><em>{company}</em></p>',
        ]
        if job.get('goal'):
            row_parts.append(f'<p>{render_html_text("Goal: " + job["goal"])}</p>')
        if job.get('value'):
            row_parts.append(f'<p>{render_html_text("Value: " + job["value"])}</p>')
        row_parts.append('<p><em>My Contribution:</em></p>')
        row_parts.append('<ul class="bullet-list">')
        row_parts.extend(f'<li>{render_html_text(bullet)}</li>' for bullet in job['contributions'])
        row_parts.append('</ul>')
        row_parts.append('</td>')
        row_parts.append('</tr>')
        experience_rows.append('\n'.join(row_parts))

    technical_lines = []
    for item in data['technical_experience']:
        label = html.escape(item['label'], quote=False)
        details = html.escape(item['details'], quote=False)
        style = ''
        if item.get('hanging_indent'):
            style = ' style="margin-left: 0.50in; text-indent: -0.50in;"'
        technical_lines.append(f'<p{style}><strong>{label}:</strong> {details}</p>')

    technical_row = '\n'.join([
        '<tr class="section-row">',
        '<td></td>',
        '<td><strong>TECHNICAL EXPERIENCE</strong></td>',
        '<td>',
        '\n'.join(technical_lines),
        '</td>',
        '</tr>',
    ])

    replacements = {
        '{{TITLE}}': f'{html.escape(data["name"], quote=False)} - Resume',
        '{{NAME_ROW}}': name_row,
        '{{CONTACT_ROW}}': contact_row,
        '{{EDUCATION_ROW}}': education_row,
        '{{EXPERIENCE_HEADER_ROW}}': experience_header_row,
        '{{EXPERIENCE_ROWS}}': '\n'.join(experience_rows),
        '{{TECHNICAL_ROW}}': technical_row,
    }
    for key, value in replacements.items():
        template = template.replace(key, value)
    return template


def build_docx(data, output_path):
    doc = Document()
    set_style_fonts(doc.styles['Normal'], 'Times New Roman', 11)

    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_paragraph_format(name_para)
    name_run = name_para.add_run(data['name'])
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_para.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=0, cols=3)
    table.autofit = False
    table.allow_autofit = False
    set_table_indent(table, -0.25)

    col_widths = [Inches(0.17), Inches(1.10), Inches(5.96)]

    def add_row():
        row = table.add_row()
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width
        return row

    # Contact row
    row = add_row()
    clear_cell(row.cells[0])
    contact_cell = row.cells[1].merge(row.cells[2])
    clear_cell(contact_cell)
    set_cell_border(contact_cell, bottom={"val": "single", "sz": "12", "space": "0", "color": "000000"})
    for line in data['contact_lines']:
        p = contact_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        apply_paragraph_format(p)
        # Check if line contains a URL
        url_match = URL_RE.search(line)
        if url_match:
            url_text = url_match.group(1)
            url_href = make_url(url_text)
            # Add text before URL if any
            before = line[:url_match.start()]
            if before:
                add_text_runs(p, before)
            add_hyperlink(p, url_href, url_text)
            # Add text after URL if any
            after = line[url_match.end():]
            if after:
                add_text_runs(p, after)
        else:
            add_text_runs(p, line)

    # Education row
    row = add_row()
    clear_cell(row.cells[0])
    clear_cell(row.cells[1])
    header_p = row.cells[1].add_paragraph()
    apply_paragraph_format(header_p)
    header_run = header_p.add_run('EDUCATION')
    header_run.bold = True

    clear_cell(row.cells[2])
    edu = data['education']

    degree_p = row.cells[2].add_paragraph()
    apply_paragraph_format(degree_p)
    add_text_runs(degree_p, edu['degree'], bold=True)

    institution_p = row.cells[2].add_paragraph()
    apply_paragraph_format(institution_p)
    add_text_runs(institution_p, edu['institution'])

    major_p = row.cells[2].add_paragraph()
    apply_paragraph_format(major_p)
    add_text_runs(major_p, 'Major:', bold=True)
    add_text_runs(major_p, f" {edu['major']}")

    spec_p = row.cells[2].add_paragraph()
    apply_paragraph_format(spec_p)
    add_text_runs(spec_p, 'Specialization:', bold=True)
    add_text_runs(spec_p, f" {edu['specialization']}")

    # Experience header row
    row = add_row()
    clear_cell(row.cells[0])
    clear_cell(row.cells[1])
    exp_p = row.cells[1].add_paragraph()
    apply_paragraph_format(exp_p)
    exp_run = exp_p.add_run('EXPERIENCE')
    exp_run.bold = True
    exp_run.italic = True
    clear_cell(row.cells[2])

    # Experience rows
    for job in data['experience']:
        row = add_row()
        clear_cell(row.cells[0])
        clear_cell(row.cells[1])
        date_p = row.cells[1].add_paragraph()
        apply_paragraph_format(date_p)
        add_text_runs(date_p, normalize_date(job['date']), italic=True)

        clear_cell(row.cells[2])
        title_p = row.cells[2].add_paragraph()
        apply_paragraph_format(title_p)
        add_text_runs(title_p, f"{job['title']}, ", bold=True)
        add_text_runs(title_p, job['company'], italic=True)

        if job.get('goal'):
            goal_p = row.cells[2].add_paragraph()
            apply_paragraph_format(goal_p)
            add_text_runs(goal_p, f"Goal: {job['goal']}")

        if job.get('value'):
            value_p = row.cells[2].add_paragraph()
            apply_paragraph_format(value_p)
            add_text_runs(value_p, f"Value: {job['value']}")

        contrib_label_p = row.cells[2].add_paragraph()
        apply_paragraph_format(contrib_label_p)
        add_text_runs(contrib_label_p, 'My Contribution:', italic=True)

        for bullet in job['contributions']:
            bullet_p = row.cells[2].add_paragraph(style='List Bullet')
            apply_paragraph_format(bullet_p)
            add_text_runs(bullet_p, bullet)

    # Spacer row
    spacer_row = add_row()
    clear_cell(spacer_row.cells[0])
    clear_cell(spacer_row.cells[1])
    clear_cell(spacer_row.cells[2])

    # Technical Experience
    row = add_row()
    clear_cell(row.cells[0])
    clear_cell(row.cells[1])
    tech_header_p = row.cells[1].add_paragraph()
    apply_paragraph_format(tech_header_p)
    tech_run = tech_header_p.add_run('TECHNICAL EXPERIENCE')
    tech_run.bold = True

    clear_cell(row.cells[2])
    for item in data['technical_experience']:
        tech_p = row.cells[2].add_paragraph()
        apply_paragraph_format(tech_p)
        if item.get('hanging_indent'):
            tech_p.paragraph_format.left_indent = Inches(0.5)
            tech_p.paragraph_format.first_line_indent = Inches(-0.5)
        add_text_runs(tech_p, f"{item['label']}:", bold=True)
        add_text_runs(tech_p, f" {item['details']}")

    # Spacer row
    spacer_row = add_row()
    clear_cell(spacer_row.cells[0])
    clear_cell(spacer_row.cells[1])
    clear_cell(spacer_row.cells[2])

    doc.save(output_path)


def main():
    data = load_manifest(MANIFEST_PATH)

    print(f"Reading {MANIFEST_PATH}...")

    print("Rendering DOCX...")
    build_docx(data, DOCX_PATH)
    print(f"\u2713 Created {DOCX_PATH}")

    print("Rendering HTML...")
    html_content = render_html(data, include_wrapper=True)
    with open(HTML_PATH, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"\u2713 Created {HTML_PATH}")

    print("Rendering Markdown...")
    md_content = render_html(data, include_wrapper=False)
    with open(MD_PATH, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f"\u2713 Created {MD_PATH}")

    print("Rendering PDF...")
    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(PDF_PATH)
        print(f"\u2713 Created {PDF_PATH}")
    except Exception as exc:
        print(f"\u26a0 PDF conversion failed: {exc}")
        print("   You can print the HTML file to PDF from your browser for better results.")

    print("\nRendering complete!")


if __name__ == '__main__':
    main()
